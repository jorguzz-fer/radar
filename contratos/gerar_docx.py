from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Cm(3)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3)
section.right_margin = Cm(2.5)

# Styles
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(4)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.space_after = Pt(20)

def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        rest = p.add_run(text)
        rest.font.name = 'Times New Roman'
        rest.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        rest = p.add_run(text)
        rest.font.name = 'Times New Roman'
        rest.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    return p

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)

# ─── CABEÇALHO ──────────────────────────────────────────────────────────────
add_title(doc, 'CONTRATO DE PRESTAÇÃO DE SERVIÇOS DE CONSULTORIA EM')
add_title(doc, 'TECNOLOGIA DA INFORMAÇÃO E MARKETING DIGITAL')
add_subtitle(doc, 'Proposta Comercial de Referência nº 001/2026  ·  Data: 27/04/2026')

add_hr(doc)

add_para(doc,
    'As partes abaixo qualificadas celebram o presente Contrato de Prestação de Serviços de '
    'Consultoria em Tecnologia da Informação e Marketing Digital ("Contrato"), que se regerá '
    'pelas cláusulas e condições seguintes, bem como pela legislação vigente aplicável, em especial '
    'o Código Civil Brasileiro (Lei nº 10.406/2002) e a Lei Geral de Proteção de Dados '
    '(Lei nº 13.709/2018).')

# ─── QUADRO DAS PARTES ──────────────────────────────────────────────────────
doc.add_paragraph()

# Tabela partes
tbl = doc.add_table(rows=1, cols=1)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

def add_party_row(tbl, label, value, shaded=False):
    row = tbl.add_row()
    cell = row.cells[0]
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run_l = p.add_run(label + '  ')
    run_l.bold = True
    run_l.font.size = Pt(10)
    run_l.font.name = 'Times New Roman'
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)
    run_v.font.name = 'Times New Roman'
    if shaded:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0')
        tcPr.append(shd)

def add_section_header_row(tbl, text):
    row = tbl.add_row()
    cell = row.cells[0]
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E0E0E0')
    tcPr.append(shd)

# remove first empty row
tbl.rows[0].cells[0].paragraphs[0].clear()

add_section_header_row(tbl, 'CONTRATADA (Prestadora de Serviços)')
add_party_row(tbl, 'Razão Social:', 'Tudo Mudou Marketing e Consultoria')
add_party_row(tbl, 'CNPJ:', '40.790.711/0001-21')
add_party_row(tbl, 'Endereço:', 'Av. Santo Antônio, 1453 — Vila Osasco')
add_party_row(tbl, 'CEP:', '06083-210 — Osasco / SP')
add_party_row(tbl, 'Chave PIX (CNPJ):', '40.790.711/0001-21')
add_party_row(tbl, 'Representante:', '________________________________________')
add_party_row(tbl, 'Cargo:', '________________________________________')
add_party_row(tbl, 'CPF:', '________________________________________')

add_section_header_row(tbl, 'CONTRATANTE (Cliente)')
add_party_row(tbl, 'Razão Social:', 'Autron Automação Indústria e Comércio Ltda')
add_party_row(tbl, 'Nome Fantasia:', 'Autron')
add_party_row(tbl, 'CNPJ:', '72.932.718/0001-27')
add_party_row(tbl, 'Endereço:', 'R. Caetés, 601 — Térreo — Perdizes')
add_party_row(tbl, 'CEP:', '05.016-081 — São Paulo / SP')
add_party_row(tbl, 'Telefone:', '(11) 4810-7979')
add_party_row(tbl, 'Representante:', '________________________________________')
add_party_row(tbl, 'Cargo:', '________________________________________')
add_party_row(tbl, 'CPF:', '________________________________________')

doc.add_paragraph()

# ─── CLÁUSULAS ──────────────────────────────────────────────────────────────

add_heading(doc, 'CLÁUSULA 1 — OBJETO')
add_para(doc, ' O presente Contrato tem por objeto a prestação de serviços de desenvolvimento, '
    'implantação, suporte e manutenção de plataforma de Business Intelligence (BI) para a CONTRATANTE, '
    'conforme escopo técnico detalhado na Proposta Comercial nº 001/2026, incorporada a este instrumento '
    'como Anexo I.', bold_prefix='1.1')
add_para(doc, ' A plataforma consolida dados provenientes de 7 (sete) fontes distintas — Protheus ERP, '
    'CRM Ploomes e Budget/Metas — em um único ambiente de análise, composto por 8 (oito) módulos de '
    'análise operacional e financeira.', bold_prefix='1.2')
add_para(doc, ' O projeto encontra-se em estágio de produção, devidamente validado junto ao usuário '
    'final da CONTRATANTE, tendo percorrido as fases: (i) Discovery & Modelagem; (ii) Desenvolvimento '
    'Core; (iii) Módulos & Refino; e (iv) Deploy & Validação.', bold_prefix='1.3')

add_heading(doc, 'CLÁUSULA 2 — ESCOPO DOS SERVIÇOS')
add_para(doc, ' Os serviços contratados compreendem:', bold_prefix='2.1')
add_bullet(doc, 'Plataforma BI completa com os 8 (oito) módulos: (1) Visão Geral, (2) Entrada de Pedidos, '
    '(3) Disponibilidade, (4) Previsão de Entrega, (5) Estoque & SC/OP, (6) Faturamento, '
    '(7) Previsão de Faturamento e (8) Ploomes × Protheus;')
add_bullet(doc, 'Integração das 7 (sete) fontes de dados: entrada_pedido.xlsx, followup.xlsx, mata010.xlsx, '
    'faturamento.xlsx, sciozvs0.csv (Protheus), Metas_Budget_2026.xlsx (Budget) e Ganhas.xlsx (Ploomes);')
add_bullet(doc, 'Deploy e configuração em VPS com Coolify, Docker e integração GitHub privado;')
add_bullet(doc, 'Autenticação e segurança: HTTPS com SSL Let\'s Encrypt, variáveis de ambiente;')
add_bullet(doc, 'Iterações e ajustes com o usuário em produção durante o período contratual;')
add_bullet(doc, 'Suporte e manutenção evolutiva por 12 (doze) meses, incluindo correções de bugs;')
add_bullet(doc, 'Entrega do código-fonte versionado em repositório privado da CONTRATANTE.')
add_para(doc, ' A stack tecnológica é composta por Python 3.11+, Pandas, NumPy, Plotly, Docker, Coolify, '
    'GitHub e Let\'s Encrypt — sem lock-in de fornecedor proprietário.', bold_prefix='2.2')
add_para(doc, ' Não estão incluídos no escopo: (a) integrações além das 7 fontes previstas; '
    '(b) hospedagem, VPS, domínios e SSL (responsabilidade da CONTRATANTE); (c) servidores e equipamentos '
    'físicos; (d) novos módulos não previstos; (e) migração para outras stacks; (f) treinamento formal '
    '(cotado sob demanda).', bold_prefix='2.3')

add_heading(doc, 'CLÁUSULA 3 — PRAZO DE VIGÊNCIA')
add_para(doc, ' O presente Contrato vigorará pelo prazo de 12 (doze) meses, contados da data de sua '
    'assinatura, podendo ser renovado por igual período mediante acordo escrito com antecedência mínima '
    'de 30 (trinta) dias.', bold_prefix='3.1')
add_para(doc, ' O prazo abrange a manutenção, suporte e ajustes evolutivos da plataforma já entregue '
    'e em produção.', bold_prefix='3.2')
add_para(doc, ' Serviços adicionais solicitados após o encerramento do prazo serão objeto de novo '
    'instrumento ou aditivo.', bold_prefix='3.3')

add_heading(doc, 'CLÁUSULA 4 — VALOR E CONDIÇÕES DE PAGAMENTO')
add_para(doc, ' As partes elegem uma das modalidades abaixo, conforme opção assinalada na assinatura:', bold_prefix='4.1')

# Tabela de valores
tbl2 = doc.add_table(rows=5, cols=3)
tbl2.style = 'Table Grid'
headers = ['Parcela / Evento', 'Modalidade Parcelada', 'Modalidade à Vista']
for i, h in enumerate(headers):
    cell = tbl2.rows[0].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '000000')
    tcPr.append(shd)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows_data = [
    ('Setup Inicial (na assinatura)', 'R$ 2.500,00', '— incluso —'),
    ('Mensalidades', '11 × R$ 2.000,00', '— incluso —'),
    ('Desconto', '—', '15% (− R$ 3.675,00)'),
    ('Total do Contrato', 'R$ 24.500,00', 'R$ 20.825,00'),
]
for i, (a, b, c) in enumerate(rows_data):
    row = tbl2.rows[i+1]
    for j, val in enumerate([a, b, c]):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        if i == 3:
            run.bold = True

doc.add_paragraph()
p_choice = doc.add_paragraph()
p_choice.add_run('( )  Modalidade Parcelada          ( )  Modalidade à Vista').font.size = Pt(11)

add_para(doc, ' Pagamento via PIX — Favorecido: Tudo Mudou Marketing e Consultoria  |  '
    'Chave PIX (CNPJ): 40.790.711/0001-21', bold_prefix='4.2')
add_para(doc, ' Na modalidade parcelada, o setup inicial vence na assinatura e as parcelas '
    'mensais no mesmo dia dos meses seguintes.', bold_prefix='4.3')
add_para(doc, ' O atraso sujeitará a CONTRATANTE a multa moratória de 2% sobre o valor em atraso, '
    'juros de 1% ao mês e correção pelo IGPM/FGV.', bold_prefix='4.4')
add_para(doc, ' Inadimplência superior a 30 (trinta) dias faculta à CONTRATADA suspender o acesso '
    'à plataforma até a regularização.', bold_prefix='4.5')
add_para(doc, ' Em caso de renovação, os valores serão reajustados pelo IGPM/FGV acumulado nos '
    '12 meses anteriores à data de aniversário.', bold_prefix='4.6')

add_heading(doc, 'CLÁUSULA 5 — OBRIGAÇÕES DA CONTRATADA')
add_para(doc, 'A CONTRATADA obriga-se a:')
add_bullet(doc, 'Prestar os serviços com zelo, técnica e qualidade;')
add_bullet(doc, 'Manter canal de suporte responsivo em dias úteis (9h–18h), com resposta inicial em até 24h;')
add_bullet(doc, 'Corrigir bugs críticos em até 48 horas úteis após o reporte;')
add_bullet(doc, 'Manter código-fonte versionado e entregar acesso ao fim do Contrato;')
add_bullet(doc, 'Guardar sigilo sobre informações da CONTRATANTE conforme Cláusula 8;')
add_bullet(doc, 'Comunicar manutenções programadas com antecedência mínima de 5 dias úteis;')
add_bullet(doc, 'Não subcontratar sem prévia autorização escrita da CONTRATANTE.')

add_heading(doc, 'CLÁUSULA 6 — OBRIGAÇÕES DA CONTRATANTE')
add_para(doc, 'A CONTRATANTE obriga-se a:')
add_bullet(doc, 'Efetuar os pagamentos nas datas e formas pactuadas;')
add_bullet(doc, 'Fornecer os arquivos de dados nos formatos acordados (XLSX, CSV);')
add_bullet(doc, 'Contratar e manter a infraestrutura de hospedagem (VPS, domínio, SSL, internet);')
add_bullet(doc, 'Designar um interlocutor técnico para acompanhamento e aprovação de entregas;')
add_bullet(doc, 'Comunicar mudanças nas estruturas dos arquivos-fonte com antecedência de 5 dias úteis;')
add_bullet(doc, 'Não ceder ou disponibilizar a plataforma a terceiros sem autorização escrita;')
add_bullet(doc, 'Zelar pela segurança das credenciais de acesso.')

add_heading(doc, 'CLÁUSULA 7 — PROPRIEDADE INTELECTUAL E CÓDIGO-FONTE')
add_para(doc, ' O código-fonte desenvolvido será de propriedade da CONTRATANTE após o cumprimento '
    'integral das obrigações financeiras, entregue versionado em repositório privado.', bold_prefix='7.1')
add_para(doc, ' Durante a vigência, o código é de uso exclusivo da CONTRATANTE para fins previstos '
    'neste instrumento, vedada cessão ou sublicenciamento.', bold_prefix='7.2')
add_para(doc, ' A CONTRATADA reserva-se o direito de reutilizar componentes genéricos em outros '
    'projetos, desde que não exponha dados ou regras de negócio da CONTRATANTE.', bold_prefix='7.3')
add_para(doc, ' Os dados processados são de titularidade exclusiva da CONTRATANTE, armazenados na '
    'infraestrutura própria da CONTRATANTE, sem transmissão a terceiros.', bold_prefix='7.4')

add_heading(doc, 'CLÁUSULA 8 — CONFIDENCIALIDADE E LGPD')
add_para(doc, ' As partes mantêm sigilo sobre todas as informações técnicas, comerciais, financeiras '
    'e operacionais pelo período de 5 (cinco) anos após o término do Contrato.', bold_prefix='8.1')
add_para(doc, ' A CONTRATADA não divulgará nem utilizará dados da CONTRATANTE para fins próprios.', bold_prefix='8.2')
add_para(doc, ' As partes cumprem a Lei Geral de Proteção de Dados (Lei nº 13.709/2018 — LGPD) '
    'e adotam medidas técnicas adequadas de proteção de dados pessoais.', bold_prefix='8.3')
add_para(doc, ' A CONTRATADA atua como operadora de dados, processando conforme instruções '
    'da CONTRATANTE (controladora), sem finalidade autônoma.', bold_prefix='8.4')

add_heading(doc, 'CLÁUSULA 9 — RESCISÃO')
add_para(doc, ' Este Contrato poderá ser rescindido: (a) por mútuo acordo, com aviso prévio de '
    '30 dias; (b) por inadimplência, após notificação e prazo de 10 dias úteis para regularização; '
    '(c) imediatamente, por descumprimento grave das obrigações de confidencialidade.', bold_prefix='9.1')
add_para(doc, ' Rescisão antecipada sem justa causa pela CONTRATANTE: multa de 20% do valor '
    'remanescente das parcelas vincendas.', bold_prefix='9.2')
add_para(doc, ' Rescisão antecipada sem justa causa pela CONTRATADA: multa de 20% das parcelas '
    'vincendas, mais entrega do código-fonte atualizado.', bold_prefix='9.3')
add_para(doc, ' A rescisão não desobriga as partes das obrigações anteriores à sua data efetiva, '
    'nem das obrigações de confidencialidade.', bold_prefix='9.4')

add_heading(doc, 'CLÁUSULA 10 — LIMITAÇÃO DE RESPONSABILIDADE')
add_para(doc, ' A CONTRATADA não responde por falhas decorrentes da infraestrutura de hospedagem '
    'sob responsabilidade da CONTRATANTE.', bold_prefix='10.1')
add_para(doc, ' A CONTRATADA não responde por imprecisões causadas por dados-fonte incorretos '
    'ou fornecidos fora do formato acordado.', bold_prefix='10.2')
add_para(doc, ' A responsabilidade total da CONTRATADA fica limitada ao valor pago nos '
    '3 (três) meses anteriores ao evento danoso.', bold_prefix='10.3')

add_heading(doc, 'CLÁUSULA 11 — DISPOSIÇÕES GERAIS')
add_para(doc, ' Este Contrato substitui quaisquer negociações ou acordos anteriores em relação '
    'ao seu objeto.', bold_prefix='11.1')
add_para(doc, ' Alterações somente produzem efeito se formalizadas por escrito e assinadas '
    'por ambas as partes.', bold_prefix='11.2')
add_para(doc, ' Tolerância a descumprimento não implica novação, renúncia ou alteração '
    'das condições.', bold_prefix='11.3')
add_para(doc, ' Cláusulas inválidas não invalidam o restante do instrumento.', bold_prefix='11.4')
add_para(doc, ' E-mail com confirmação de recebimento é aceito como meio válido de '
    'notificação.', bold_prefix='11.5')

add_heading(doc, 'CLÁUSULA 12 — FORO')
add_para(doc, ' As partes elegem o Foro da Comarca de Osasco — Estado de São Paulo para '
    'dirimir quaisquer controvérsias, renunciando a qualquer outro.', bold_prefix='12.1')

add_hr(doc)

p_local = doc.add_paragraph()
p_local.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_local.add_run('Por estarem assim justas e contratadas, as partes assinam o presente instrumento '
    'em 2 (duas) vias de igual teor, juntamente com as 2 (duas) testemunhas abaixo.')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_d = p_date.add_run('\nOsasco / SP, ______ de __________________ de 2026.')
run_d.bold = True
run_d.font.name = 'Times New Roman'
run_d.font.size = Pt(12)

# Assinaturas
def add_sig_table(doc, left_label, left_details, right_label, right_details):
    doc.add_paragraph()
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = 'Table Grid'
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBdr = OxmlElement('w:tcBdr')
            for side in ['top', 'left', 'bottom', 'right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                tcBdr.append(el)
            tcPr.append(tcBdr)

    def set_cell(row_idx, col_idx, text, bold=False, center=True, border_top=False):
        cell = tbl.rows[row_idx].cells[col_idx]
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        if border_top:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBdr = OxmlElement('w:tcBdr')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '6')
            top.set(qn('w:space'), '1')
            top.set(qn('w:color'), '000000')
            tcBdr.append(top)
            tcPr.append(tcBdr)

    # row 0: empty (space for signature)
    set_cell(0, 0, '\n\n')
    set_cell(0, 1, '\n\n')
    # row 1: line + name
    set_cell(1, 0, '_______________________________________________', border_top=False)
    set_cell(1, 1, '_______________________________________________', border_top=False)
    # row 2: details
    set_cell(2, 0, left_details, bold=False)
    set_cell(2, 1, right_details, bold=False)

add_sig_table(doc,
    'CONTRATADA',
    'Tudo Mudou Marketing e Consultoria\nCNPJ: 40.790.711/0001-21\nRepresentante: ________________________\nCPF: _____________________',
    'CONTRATANTE',
    'Autron Automação Indústria e Comércio Ltda\nCNPJ: 72.932.718/0001-27\nRepresentante: ________________________\nCPF: _____________________'
)

doc.add_paragraph()
p_wit = doc.add_paragraph()
p_wit.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_w = p_wit.add_run('Testemunhas:')
run_w.bold = True
run_w.font.size = Pt(11)
run_w.font.name = 'Times New Roman'

add_sig_table(doc,
    'Testemunha 1',
    'Nome: ________________________________\nCPF: _____________________',
    'Testemunha 2',
    'Nome: ________________________________\nCPF: _____________________'
)

add_hr(doc)
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = p_footer.add_run(
    'Contrato nº 001/2026  ·  Tudo Mudou Marketing e Consultoria × Autron Automação Indústria e Comércio Ltda\n'
    'Gerado em 08/05/2026')
run_f.font.size = Pt(8)
run_f.font.name = 'Times New Roman'
run_f.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save('/home/user/radar/contratos/contrato_autron.docx')
print('DOCX gerado com sucesso.')
